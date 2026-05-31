import json
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.models.project import Project, ProjectMember
from app.models.room import Room
from app.models.message import Message, MessageType, ConsensusStatus
from app.models.vote import Vote
from app.models.instruction import Instruction
from app.models.decision import Decision
from app.models.user import User
from jinja2 import Environment, BaseLoader


MARKDOWN_TEMPLATE = """# {{ project.name }} — SciFig Collaboration Export

**Description:** {{ project.description or 'N/A' }}
**Exported:** {{ now }}

---

## Discussion Log

{% for msg in messages %}
### [{{ msg.message_type | upper }}] {{ msg.user_display_name }} ({{ msg.role }}) — {{ msg.created_at }}
{% if msg.change_category %}**Category:** {{ msg.change_category }} | **Status:** {{ msg.consensus_status }}{% endif %}

{{ msg.content }}

{% if msg.votes %}**Votes:** ✅ {{ msg.upvotes }} / ❌ {{ msg.downvotes }}{% endif %}

---
{% endfor %}

## Approved Instructions

{% for instr in approved_instructions %}
- **{{ instr.instruction_type | upper }}** → Target: `{{ instr.target }}` | Action: `{{ instr.action }}`
  Payload: `{{ instr.payload }}`
{% endfor %}

## Decisions

{% for dec in decisions %}
- [{{ dec.status | upper }}] {{ dec.description }} *({{ dec.decided_at or 'Pending' }})*
{% endfor %}
"""

HTML_TEMPLATE = """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body { font-family: Arial, sans-serif; max-width: 900px; margin: 40px auto; color: #1a1a2e; }
h1 { color: #6366f1; border-bottom: 3px solid #6366f1; padding-bottom: 10px; }
h2 { color: #4f46e5; margin-top: 30px; }
.message { border-left: 4px solid #6366f1; padding: 12px; margin: 16px 0; background: #f8f9ff; border-radius: 4px; }
.meta { font-size: 12px; color: #666; margin-bottom: 8px; }
.badge { display: inline-block; padding: 2px 8px; border-radius: 12px; font-size: 11px; font-weight: bold; margin-right: 6px; }
.suggestion { background: #e0e0ff; color: #4f46e5; }
.approval { background: #dcfce7; color: #16a34a; }
.objection { background: #fee2e2; color: #dc2626; }
.feedback { background: #fef3c7; color: #d97706; }
.accepted { background: #dcfce7; color: #16a34a; }
.rejected { background: #fee2e2; color: #dc2626; }
.pending_vote { background: #fef3c7; color: #d97706; }
table { width: 100%; border-collapse: collapse; margin: 16px 0; }
th { background: #6366f1; color: white; padding: 10px; text-align: left; }
td { padding: 10px; border-bottom: 1px solid #e5e7eb; }
</style>
</head>
<body>
<h1>{{ project.name }} — SciFig Collaboration Export</h1>
<p><strong>Description:</strong> {{ project.description or 'N/A' }}<br>
<strong>Exported:</strong> {{ now }}</p>

<h2>Discussion Log</h2>
{% for msg in messages %}
<div class="message">
  <div class="meta">
    <span class="badge {{ msg.message_type }}">{{ msg.message_type | upper }}</span>
    <strong>{{ msg.user_display_name }}</strong> ({{ msg.role }}) — {{ msg.created_at }}
    {% if msg.consensus_status %}<span class="badge {{ msg.consensus_status }}">{{ msg.consensus_status | replace('_',' ') | upper }}</span>{% endif %}
  </div>
  <p>{{ msg.content }}</p>
  {% if msg.upvotes or msg.downvotes %}<small>✅ {{ msg.upvotes }} upvotes / ❌ {{ msg.downvotes }} downvotes</small>{% endif %}
</div>
{% endfor %}

<h2>Approved Instructions</h2>
<table>
<tr><th>Type</th><th>Target</th><th>Action</th><th>Payload</th></tr>
{% for instr in approved_instructions %}
<tr><td>{{ instr.instruction_type }}</td><td>{{ instr.target }}</td><td>{{ instr.action }}</td><td>{{ instr.payload }}</td></tr>
{% endfor %}
</table>

<h2>Decisions</h2>
<table>
<tr><th>Status</th><th>Description</th><th>Date</th></tr>
{% for dec in decisions %}
<tr><td><span class="badge {{ dec.status }}">{{ dec.status | upper }}</span></td><td>{{ dec.description }}</td><td>{{ dec.decided_at or 'Pending' }}</td></tr>
{% endfor %}
</table>
</body>
</html>"""


async def _build_export_data(project_id: str, db: AsyncSession) -> dict:
    proj_result = await db.execute(select(Project).where(Project.id == project_id))
    project = proj_result.scalar_one_or_none()

    rooms_result = await db.execute(select(Room).where(Room.project_id == project_id))
    rooms = rooms_result.scalars().all()
    room_ids = [r.id for r in rooms]

    messages_data = []
    for room_id in room_ids:
        msgs_result = await db.execute(
            select(Message, User)
            .outerjoin(User, User.id == Message.user_id)
            .where(Message.room_id == room_id)
            .order_by(Message.created_at)
        )
        for msg, user in msgs_result.all():
            votes_result = await db.execute(select(Vote).where(Vote.message_id == msg.id))
            votes = votes_result.scalars().all()
            upvotes = sum(1 for v in votes if v.vote_type == "upvote")
            downvotes = sum(1 for v in votes if v.vote_type == "downvote")
            messages_data.append({
                "id": msg.id,
                "room_id": room_id,
                "user_display_name": user.display_name if user else "Deleted User",
                "role": "unknown",
                "content": msg.content,
                "message_type": msg.message_type,
                "change_category": msg.change_category,
                "consensus_status": msg.consensus_status,
                "upvotes": upvotes,
                "downvotes": downvotes,
                "votes": len(votes),
                "created_at": msg.created_at.isoformat(),
            })

    instrs_result = await db.execute(select(Instruction).where(Instruction.project_id == project_id, Instruction.status == "applied"))
    approved_instructions = [{"instruction_type": i.instruction_type, "target": i.target, "action": i.action, "payload": i.payload} for i in instrs_result.scalars().all()]

    decs_result = await db.execute(select(Decision).where(Decision.project_id == project_id))
    decisions = [{"description": d.description, "status": d.status, "decided_at": d.decided_at.isoformat() if d.decided_at else None} for d in decs_result.scalars().all()]

    return {"project": project, "messages": messages_data, "approved_instructions": approved_instructions, "decisions": decisions}


async def export_json(project_id: str, db: AsyncSession) -> dict:
    data = await _build_export_data(project_id, db)
    return {
        "project": {"id": str(data["project"].id), "name": data["project"].name, "description": data["project"].description},
        "messages": data["messages"],
        "approved_instructions": data["approved_instructions"],
        "decisions": data["decisions"],
    }


async def export_markdown(project_id: str, db: AsyncSession) -> str:
    from datetime import datetime, timezone
    data = await _build_export_data(project_id, db)
    env = Environment(loader=BaseLoader())
    tmpl = env.from_string(MARKDOWN_TEMPLATE)
    return tmpl.render(project=data["project"], messages=data["messages"], approved_instructions=data["approved_instructions"], decisions=data["decisions"], now=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))


async def export_pdf(project_id: str, db: AsyncSession) -> bytes:
    from datetime import datetime, timezone
    import weasyprint
    data = await _build_export_data(project_id, db)
    env = Environment(loader=BaseLoader())
    tmpl = env.from_string(HTML_TEMPLATE)
    html = tmpl.render(project=data["project"], messages=data["messages"], approved_instructions=data["approved_instructions"], decisions=data["decisions"], now=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
    return weasyprint.HTML(string=html).write_pdf()


async def export_docx(project_id: str, db: AsyncSession) -> bytes:
    from io import BytesIO
    from datetime import datetime, timezone
    from docx import Document
    from docx.shared import Pt, RGBColor
    data = await _build_export_data(project_id, db)
    doc = Document()
    doc.add_heading(f"{data['project'].name} — SciFig Export", 0)
    doc.add_paragraph(f"Description: {data['project'].description or 'N/A'}")
    doc.add_paragraph(f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_heading("Discussion Log", 1)
    for msg in data["messages"]:
        p = doc.add_paragraph()
        p.add_run(f"[{msg['message_type'].upper()}] ").bold = True
        p.add_run(f"{msg['user_display_name']} — {msg['created_at']}\n")
        p.add_run(msg["content"])
        if msg["votes"]:
            p.add_run(f"\n✅ {msg['upvotes']} / ❌ {msg['downvotes']}")
    doc.add_heading("Approved Instructions", 1)
    for instr in data["approved_instructions"]:
        doc.add_paragraph(f"• {instr['instruction_type'].upper()} → {instr['target']} : {instr['action']}", style="List Bullet")
    doc.add_heading("Decisions", 1)
    for dec in data["decisions"]:
        doc.add_paragraph(f"[{dec['status'].upper()}] {dec['description']}", style="List Bullet")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
